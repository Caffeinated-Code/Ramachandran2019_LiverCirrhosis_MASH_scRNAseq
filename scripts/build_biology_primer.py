from pathlib import Path
from textwrap import wrap

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
ASSETS = DOCS / "assets" / "primer"
OUT = DOCS / "biology_primer_liver_fibrosis.docx"

W, H = 1400, 760
INK = "#1F2933"
MUTED = "#52606D"
BLUE = "#2F80ED"
GREEN = "#219653"
ORANGE = "#F2994A"
RED = "#EB5757"
PURPLE = "#9B51E0"
GRAY = "#E5E7EB"
LIGHT_BLUE = "#EAF3FF"
LIGHT_GREEN = "#EAF7EF"
LIGHT_ORANGE = "#FFF3E6"
LIGHT_RED = "#FDECEC"
LIGHT_PURPLE = "#F3EAFE"


def font(size, bold=False):
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/Library/Fonts/Arial Bold.ttf" if bold else "/Library/Fonts/Arial.ttf",
        "/System/Library/Fonts/Helvetica.ttc",
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size)
        except Exception:
            continue
    return ImageFont.load_default()


def rounded(draw, box, fill, outline=None, width=2, radius=28):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def text(draw, xy, value, size=34, fill=INK, bold=False, anchor=None):
    draw.text(xy, value, fill=fill, font=font(size, bold), anchor=anchor)


def wrapped(draw, xy, value, width_chars=28, size=30, fill=INK, bold=False, line_gap=8):
    x, y = xy
    for line in wrap(value, width_chars):
        draw.text((x, y), line, fill=fill, font=font(size, bold))
        y += size + line_gap
    return y


def arrow(draw, start, end, fill=MUTED, width=5):
    draw.line([start, end], fill=fill, width=width)
    x1, y1 = start
    x2, y2 = end
    if x2 >= x1:
        head = [(x2, y2), (x2 - 22, y2 - 12), (x2 - 22, y2 + 12)]
    else:
        head = [(x2, y2), (x2 + 22, y2 - 12), (x2 + 22, y2 + 12)]
    draw.polygon(head, fill=fill)


def save_progression():
    img = Image.new("RGB", (W, H), "white")
    d = ImageDraw.Draw(img)
    text(d, (70, 55), "How liver scarring builds over time", 44, bold=True)
    text(d, (70, 112), "Repeated injury can turn normal repair into long-lasting scar tissue.", 28, MUTED)

    cards = [
        ("Healthy liver", "Flexible tissue carries blood and performs metabolism.", LIGHT_GREEN, GREEN),
        ("Metabolic stress", "Fat, insulin resistance, and inflammation injure liver cells.", LIGHT_ORANGE, ORANGE),
        ("Fibrosis", "Repair cells lay down collagen, like a patch that becomes too thick.", LIGHT_RED, RED),
        ("Cirrhosis", "Scar changes the liver's structure and blood flow.", LIGHT_PURPLE, PURPLE),
    ]
    x = 70
    y = 235
    for i, (title, body, fill, accent) in enumerate(cards):
        rounded(d, (x, y, x + 285, y + 275), fill, accent, 4)
        d.ellipse((x + 95, y + 30, x + 190, y + 125), fill="white", outline=accent, width=5)
        if i == 0:
            d.arc((x + 105, y + 45, x + 180, y + 110), 210, 510, fill=accent, width=5)
        elif i == 1:
            d.ellipse((x + 112, y + 52, x + 138, y + 78), fill=accent)
            d.ellipse((x + 148, y + 85, x + 173, y + 110), fill=accent)
        elif i == 2:
            for off in [0, 20, 40]:
                d.line((x + 112 + off, y + 50, x + 165 + off, y + 110), fill=accent, width=5)
        else:
            d.line((x + 115, y + 50, x + 175, y + 110), fill=accent, width=6)
            d.line((x + 175, y + 50, x + 115, y + 110), fill=accent, width=6)
            d.arc((x + 112, y + 52, x + 178, y + 112), 15, 345, fill=accent, width=4)
        text(d, (x + 28, y + 142), title, 28, INK, True)
        wrapped(d, (x + 28, y + 184), body, 27, 20, INK, line_gap=5)
        if i < len(cards) - 1:
            arrow(d, (x + 300, y + 138), (x + 365, y + 138))
        x += 335

    path = ASSETS / "liver_fibrosis_progression.png"
    img.save(path)
    return path


def save_niche():
    img = Image.new("RGB", (W, H), "white")
    d = ImageDraw.Draw(img)
    text(d, (70, 55), "The fibrotic niche: several cell types talking at once", 42, bold=True)
    text(d, (70, 112), "Single-cell data helps separate the voices inside scarred tissue.", 28, MUTED)

    d.ellipse((480, 190, 920, 610), fill="#FFF7ED", outline=ORANGE, width=6)
    for i in range(9):
        x = 540 + (i % 3) * 105
        y = 250 + (i // 3) * 80
        d.line((x, y, x + 120, y + 65), fill=RED, width=8)
    text(d, (700, 405), "Scar matrix", 34, RED, True, anchor="mm")

    cells = [
        ((120, 245, 380, 365), "Stellate / myofibroblast", "builds collagen scar", LIGHT_RED, RED),
        ((1010, 245, 1280, 365), "Macrophage", "immune signaler", LIGHT_BLUE, BLUE),
        ((120, 485, 380, 605), "Endothelial cell", "blood vessel lining", LIGHT_GREEN, GREEN),
        ((1010, 485, 1280, 605), "Injured liver cell", "stress signal source", LIGHT_ORANGE, ORANGE),
    ]
    for box, title, body, fill, accent in cells:
        rounded(d, box, fill, accent, 4)
        text(d, (box[0] + 24, box[1] + 24), title, 27, INK, True)
        wrapped(d, (box[0] + 24, box[1] + 65), body, 21, 24, INK)

    arrow(d, (380, 305), (510, 340), RED)
    arrow(d, (1010, 305), (900, 340), BLUE)
    arrow(d, (380, 545), (510, 485), GREEN)
    arrow(d, (1010, 545), (900, 485), ORANGE)
    text(d, (70, 665), "Key idea: fibrosis is a neighborhood problem, not one rogue gene in one cell.", 30, INK, True)
    path = ASSETS / "fibrotic_niche_cells.png"
    img.save(path)
    return path


def save_scrna():
    img = Image.new("RGB", (W, H), "white")
    d = ImageDraw.Draw(img)
    text(d, (70, 55), "What single-cell RNA-seq adds", 44, bold=True)
    text(d, (70, 112), "Instead of blending a tissue sample, it reads gene activity cell by cell.", 28, MUTED)

    steps = [
        ("Tissue", "many mixed cells", LIGHT_ORANGE, ORANGE),
        ("Single cells", "one barcode per cell", LIGHT_BLUE, BLUE),
        ("Gene readout", "which genes are on?", LIGHT_GREEN, GREEN),
        ("Map", "cell types and disease states", LIGHT_PURPLE, PURPLE),
        ("Prioritize", "biomarkers and targets", LIGHT_RED, RED),
    ]
    x = 70
    y = 255
    for i, (title, body, fill, accent) in enumerate(steps):
        rounded(d, (x, y, x + 225, y + 230), fill, accent, 4)
        text(d, (x + 28, y + 28), title, 29, INK, True)
        wrapped(d, (x + 28, y + 78), body, 18, 23, INK)
        if title == "Tissue":
            for cx, cy, c in [(x+75,y+165,RED),(x+120,y+145,BLUE),(x+155,y+175,GREEN)]:
                d.ellipse((cx-24, cy-24, cx+24, cy+24), fill=c)
        elif title == "Single cells":
            for j, c in enumerate([RED, BLUE, GREEN, PURPLE]):
                d.ellipse((x+55+j*38, y+155, x+82+j*38, y+182), fill=c)
        elif title == "Gene readout":
            for j, h in enumerate([45, 85, 60, 105]):
                d.rectangle((x+50+j*35, y+190-h, x+72+j*35, y+190), fill=accent)
        elif title == "Map":
            for cx, cy, c in [(x+70,y+145,RED),(x+96,y+170,RED),(x+155,y+145,BLUE),(x+175,y+175,BLUE),(x+125,y+195,GREEN)]:
                d.ellipse((cx-10, cy-10, cx+10, cy+10), fill=c)
        else:
            d.polygon([(x+70,y+145),(x+170,y+145),(x+140,y+195),(x+100,y+195)], fill=accent)
        if i < len(steps) - 1:
            arrow(d, (x + 235, y + 115), (x + 275, y + 115))
        x += 265

    text(d, (70, 635), "Why it matters: a gene can look important only in one cell state, even if it is hidden in bulk tissue.", 28, INK)
    path = ASSETS / "single_cell_workflow.png"
    img.save(path)
    return path


def save_prioritization():
    img = Image.new("RGB", (W, H), "white")
    d = ImageDraw.Draw(img)
    text(d, (70, 55), "From thousands of genes to a short, useful list", 42, bold=True)
    text(d, (70, 112), "A good target is not just statistically different. It has to make biological and practical sense.", 27, MUTED)

    layers = [
        ("Disease signal", "changes with fibrosis", BLUE),
        ("Cell specificity", "points to the right compartment", GREEN),
        ("Mechanism", "fits fibrosis biology", ORANGE),
        ("Validation", "seen in outside data", PURPLE),
        ("Translation", "preclinical fit", RED),
    ]
    top = 210
    left = 250
    widths = [900, 760, 610, 470, 330]
    for i, (title, body, color) in enumerate(layers):
        y = top + i * 88
        x0 = left + (900 - widths[i]) // 2
        x1 = x0 + widths[i]
        d.polygon([(x0, y), (x1, y), (x1 - 50, y + 64), (x0 + 50, y + 64)], fill=color)
        text(d, ((x0 + x1) // 2, y + 18), title, 25, "white", True, anchor="mt")
        text(d, ((x0 + x1) // 2, y + 42), body, 20, "white", False, anchor="mt")
    rounded(d, (500, 665, 900, 725), "#F8FAFC", INK, 3)
    text(d, (700, 681), "Prioritized candidates", 30, INK, True, anchor="ma")
    path = ASSETS / "target_prioritization_funnel.png"
    img.save(path)
    return path


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill.replace("#", ""))
    tc_pr.append(shd)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_font(run, size=11, color="000000", bold=False, italic=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def paragraph(doc, text_value="", style=None, after=6, before=0):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if text_value:
        run = p.add_run(text_value)
        set_font(run)
    return p


def add_heading(doc, value, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.space_before = Pt(16 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(value)
    set_font(run, size=16 if level == 1 else 13, color="2E74B5" if level <= 2 else "1F4D78", bold=True)
    return p


def add_image(doc, path, caption):
    doc.add_picture(str(path), width=Inches(6.3))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(caption)
    set_font(r, size=9, color="52606D", italic=True)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        set_font(r, size=10.5)


def add_table(doc):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for cell, value in zip(hdr, ["Cell type", "Plain-language role", "Why it matters in fibrosis"]):
        set_cell_shading(cell, "#F2F4F7")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value)
        set_font(r, size=10.5, bold=True)
    rows = [
        ("Hepatocyte", "Main working cell of the liver.", "Metabolic stress and injury start many inflammatory signals."),
        ("Hepatic stellate cell", "Normally stores vitamin A; when activated, becomes a scar builder.", "Central source of collagen and matrix remodeling."),
        ("Macrophage", "Immune cell that cleans up damage and sends instructions.", "Can support repair, but disease-associated states can push fibrosis."),
        ("Endothelial cell", "Lines blood vessels and controls traffic between blood and tissue.", "Changes in scarred liver affect immune entry, oxygen delivery, and niche signaling."),
        ("Cholangiocyte", "Lines bile ducts.", "Can expand or signal during chronic injury and ductular reaction."),
    ]
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            set_font(r, size=10)
    return table


def build_doc():
    ASSETS.mkdir(parents=True, exist_ok=True)
    images = {
        "progression": save_progression(),
        "niche": save_niche(),
        "scrna": save_scrna(),
        "prioritization": save_prioritization(),
    }

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for idx, size in [(1, 16), (2, 13), (3, 12)]:
        style = styles[f"Heading {idx}"]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string("2E74B5" if idx < 3 else "1F4D78")
        style.paragraph_format.space_before = Pt(16 if idx == 1 else 12)
        style.paragraph_format.space_after = Pt(6)

    header = section.header.paragraphs[0]
    header.text = "Biology primer | Human liver fibrosis"
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = RGBColor.from_string("52606D")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.text = "Prepared as background for single-cell target discovery"
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor.from_string("52606D")

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(8)
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("Liver Fibrosis, MASH, and the Fibrotic Niche")
    set_font(run, size=24, color="0B2545", bold=True)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    run = subtitle.add_run("A plain-language biology primer for target discovery discussions")
    set_font(run, size=13, color="52606D")

    add_heading(doc, "The short version", 1)
    paragraph(
        doc,
        "Fibrosis means scar tissue. In the liver, scar forms when injury and repair keep repeating. "
        "At first, scar can be part of healing. Over time, too much scar stiffens the liver, changes blood flow, "
        "and disrupts normal metabolism. Cirrhosis is advanced scarring where the liver's structure is deeply remodeled.",
    )
    paragraph(
        doc,
        "MASH, formerly called NASH, is a common route into fibrosis. It begins with metabolic stress, fat buildup, "
        "and inflammation. The important point for target discovery is that fibrosis is made by a community of cells, "
        "not by one cell type acting alone.",
    )
    add_image(doc, images["progression"], "Figure 1. Repeated injury can shift repair from temporary patching to long-lasting scar.")

    add_heading(doc, "The main cell players", 1)
    paragraph(
        doc,
        "A useful way to understand liver fibrosis is to picture a damaged neighborhood. Some cells send alarm signals, "
        "some cells try to clean up damage, some cells build scar, and blood-vessel cells change the local traffic pattern.",
    )
    add_table(doc)
    paragraph(
        doc,
        "In a healthy repair response, these jobs are coordinated and temporary. In chronic disease, the same repair machinery "
        "can become persistent and harmful.",
    )
    add_image(doc, images["niche"], "Figure 2. The fibrotic niche is the local scar environment where multiple cell types exchange signals.")

    add_heading(doc, "Why single-cell data is useful", 1)
    paragraph(
        doc,
        "A liver sample contains many cell types mixed together. Bulk RNA-seq is like hearing the whole room at once. "
        "Single-cell RNA-seq puts a barcode on each cell and asks which genes are active in that cell. This helps separate "
        "a macrophage signal from a stellate-cell signal or an endothelial-cell signal.",
    )
    add_image(doc, images["scrna"], "Figure 3. Single-cell RNA-seq separates mixed tissue into cell-type and cell-state signals.")

    add_heading(doc, "What the Ramachandran cirrhosis dataset puts in focus", 1)
    paragraph(
        doc,
        "The Ramachandran et al. study mapped human cirrhotic liver at single-cell resolution. It highlighted disease-associated "
        "cell populations that sit within the scarred niche: TREM2+CD9+ scar-associated macrophages, ACKR1+ and PLVAP+ "
        "scar-associated endothelial cells, and PDGFRA+ collagen-producing mesenchymal cells.",
    )
    paragraph(
        doc,
        "For this project, those findings are not treated as trivia. They are anchors for quality control and interpretation. "
        "If an analysis cannot recover the expected macrophage, endothelial, and mesenchymal disease programs, the downstream "
        "target list should be treated with skepticism.",
    )

    add_heading(doc, "Biomarker is not the same as target", 1)
    paragraph(
        doc,
        "A biomarker is a measurable sign. A therapeutic target is something we might try to change with a drug. A smoke alarm "
        "is a good sign of fire, but it is not the fire itself. In the same way, a gene can be an excellent disease marker but "
        "a poor drug target.",
    )
    add_bullets(
        doc,
        [
            "Good diagnostic biomarkers are measurable and track disease state.",
            "Good pharmacodynamic biomarkers change when a therapy affects the pathway.",
            "Good therapeutic targets are connected to disease biology, accessible to a drug, and tolerable to perturb.",
            "Good preclinical targets should have enough conservation in model systems to support animal studies.",
        ],
    )
    add_image(doc, images["prioritization"], "Figure 4. Target prioritization narrows genes through biology, validation, and translational filters.")

    add_heading(doc, "How to read candidate evidence", 1)
    paragraph(
        doc,
        "A candidate becomes more convincing when independent lines of evidence point in the same direction. For example, a gene "
        "that is increased in cirrhotic mesenchymal cells, enriched in a fibrosis pathway, detected in an external MASH dataset, "
        "and encodes a surface or secreted protein is easier to prioritize than a gene supported by one statistical test alone.",
    )
    paragraph(
        doc,
        "The reverse is also true. A candidate should be downgraded if it is broadly expressed in many healthy tissues, lacks "
        "disease specificity, is not conserved enough for model testing, or appears only in a small number of cells from one donor.",
    )

    add_heading(doc, "Vocabulary", 1)
    add_bullets(
        doc,
        [
            "Fibrosis: buildup of scar-like extracellular matrix after repeated injury.",
            "Cirrhosis: advanced fibrosis with major changes to liver structure and function.",
            "MASLD: metabolic dysfunction-associated steatotic liver disease, the updated term replacing much of the older NAFLD language.",
            "MASH: metabolic dysfunction-associated steatohepatitis, the updated term for NASH.",
            "Hepatic stellate cell: a liver cell that can become a major scar-producing myofibroblast.",
            "Macrophage: an immune cell that can clean up damage and send inflammatory or repair signals.",
            "Endothelial cell: a cell lining blood vessels that helps control vascular and immune traffic.",
            "Extracellular matrix: the structural material around cells; collagen is a major component of scar.",
            "Pseudobulk: an analysis strategy that aggregates cells by donor or sample before testing disease effects.",
        ],
    )

    add_heading(doc, "Selected sources", 1)
    sources = [
        "Ramachandran P. et al. Resolving the fibrotic niche of human liver cirrhosis at single-cell level. Nature. 2019. https://www.nature.com/articles/s41586-019-1631-3",
        "Rinella M. E. et al. A multisociety Delphi consensus statement on new fatty liver disease nomenclature. Hepatology. 2023. https://pubmed.ncbi.nlm.nih.gov/37363821/",
        "FDA Drug Trials Snapshot: Rezdiffra (resmetirom), original approval March 14, 2024. https://www.fda.gov/drugs/drug-approvals-and-databases/drug-trials-snapshots-rezdiffra",
        "Nature Reviews Gastroenterology & Hepatology. Clues to cirrhosis: a focus on fibrosis. 2019. https://www.nature.com/articles/s41575-019-0232-9",
    ]
    for source in sources:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(source)
        set_font(r, size=9.5, color="1F2933")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_doc()
